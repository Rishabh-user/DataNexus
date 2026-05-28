import asyncio
from pathlib import Path

from app.core.logging import get_logger
from app.extraction.base import BaseExtractor, ExtractionResult, TableData

logger = get_logger(__name__)


class DocxExtractor(BaseExtractor):
    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    async def extract(self, file_path: Path) -> ExtractionResult:
        return await asyncio.to_thread(self._extract_sync, file_path)

    def _extract_sync(self, file_path: Path) -> ExtractionResult:
        from docx import Document

        result = ExtractionResult()
        doc = Document(str(file_path))

        all_text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text_parts.append(para.text)

        for table in doc.tables:
            headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            rows = []
            for row in table.rows[1:]:
                rows.append([cell.text.strip() for cell in row.cells])
            if headers:
                result.tables.append(TableData(headers=headers, rows=rows))

        result.text = "\n\n".join(all_text_parts)

        props = doc.core_properties
        result.metadata = {
            "author": props.author or "",
            "title": props.title or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else "",
            "word_count": len(result.text.split()),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }
        return result
